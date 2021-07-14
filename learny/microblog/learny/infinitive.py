#! python3
# -*- coding: utf-8 -*-


def infinitive(inputtext, language):
	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt
	import os
	import random
	import spacy
	try:
		import docxprint, languageload
	except ImportError:
		from learny import docxprint, languageload
	print('-------------------GENERATE infinit form of verbs: ')



	# variables and lists used
	i = 0
	verb_list = []
	some_verbs = []
	new_text = []
	infinit_verb_list = []
	textverb = []
	keptinfinitive = []
	non_infinit_verbs = []
	new_infinit_verbs = []
	Aufgabe = {
		"Kopfzeile": "Name: 				Klasse: 				Datum:  \n ",
		"Titel": "",
		"1. Aufgabe": "Aufgabe: Hier siehst du gebeugte Verben und Verben in der Grundform. Ordne sie einander zu!\n",
		"Hinweise": "Hier sind die Wörter aus dem Text: \n ",
		"Rätselwörter": "Hier ein paar Rätselwörter aus dem Text: \n",
	}
	# create document object
	doc = docx.Document()
	nlp = languageload.language_load(language)
	docnlp = nlp(inputtext) #load to spacy
	#nlp = spacy.load(request.form['language_options'])
	doc = docx.Document()# initializing python-docx
	save_path = docxprint.docx_print(Doc= doc, save= 'infinitive')

	print(docnlp.text)
	paragraph = doc.add_paragraph()
	paragraph.text = ''
	for token in docnlp:
		if 'AUX' in token.pos_:
			print(token.lemma_)
			infinit_verb_list.append(token.lemma_)
			textverb.append(token.text)
			docxprint.docx_print(printText=str(token) + ' ', color="red", Paragraph=paragraph, Doc=doc)

		if 'VERB' in token.pos_:
			print(token.lemma_)
			infinit_verb_list.append(token.lemma_)
			textverb.append(token.text)

			docxprint.docx_print(printText=str(token) + ' ', color="red", Paragraph=paragraph, Doc=doc)
		elif 'AUX' in token.pos_:
			paragraph.add_run('')
		else:
			docxprint.docx_print(printText=str(token) + ' ', Paragraph=paragraph, Doc=doc)

	infinit_verb_list = list(set(infinit_verb_list))
	textverb = list(set(textverb))
	doc.add_page_break()
	#paragraph = doc.add_paragraph()
	docxprint.docx_print(printText=Aufgabe["1. Aufgabe"], Bold=True, Doc=doc)
	paragraph = doc.add_paragraph('')
	#random.shuffle(mix_list)
	for word in textverb:
		if word not in infinit_verb_list:
			non_infinit_verbs.append(word)
			if word == textverb[-1]:
				docxprint.docx_print(printText=word + ' ', Paragraph=paragraph, Doc=doc)
				break
			docxprint.docx_print(printText=word + ', ', Paragraph=paragraph, Doc=doc)


	print("keptinfinitive:", keptinfinitive )
	paragraph = doc.add_paragraph()
	random.shuffle(non_infinit_verbs)
	docnlp = nlp(" ".join(non_infinit_verbs)) #load to spacy
	for token in docnlp:
		docxprint.docx_print(printText=token.lemma_ + '_'*31+"\n", Paragraph=paragraph, Doc=doc)

	print('done')
	doc.save(save_path)
	return
