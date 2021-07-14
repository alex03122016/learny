#! python3 
#
#cloze test generator(all nouns): this program searches a given docx documents for nouns and substitutes them with ' _'
#
def pronounceQuiz(fileN):
	
	#starting with the clozeTest
	import random
	import pickle
	import subprocess
	import os
	import spacy
	from spacy.lang.de.examples import sentences
	from flask import request
	#python-docx module
	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt
	#lots faster loading pickle
	#with open('nltk_german_classifier_data.pickle', 'rb') as f:
	#	tagger = pickle.load(f)




	# create document object
	doc = docx.Document()

	# variables and lists used
	i = 0
	noun_list = []
	some_nouns = []
	new_text = []
	fullrawText = []
	#welcome message
	print('-------------------GENERATE pronounciation select syllaba in docx has to be in a table----')
	skip = False # it has to be here!nearly went crazy to figure it out
	nlp = spacy.load(request.form['language_options'])


	rawText = request.form['text']
	fullrawText.append(rawText)
	docnlp = nlp(rawText)
	text = rawText.split()
	#words_tagged = tagger.tag(text)
	paragraph = doc.add_paragraph()
	paragraph.text = ''

#printing the runs

	for token in docnlp:
		word = str(token) 
		if 'eu' in word:
			i = 0
			print(word.split('eu'))
			splitWord = word.split('eu')
			print(splitWord[-1])
			for part in splitWord:
				if part == splitWord[-1]:
					paragraph.add_run(part + ' ')
				elif part != splitWord[-1]:
					redWord = paragraph.add_run('eu')
					redWord.font.color.rgb = RGBColor(255,0,0) #color red every noun word
				elif part == splitWord[0]:
					paragraph.add_run(part)

				else:
					paragraph.add_run(part)

		elif 'äu' in word:
			i = 0
			print(word.split('äu'))
			splitWord = word.split('äu')
			print(splitWord[-1])
			for part in splitWord:
				if part == splitWord[-1]:
					paragraph.add_run(part + ' ')
				elif part != splitWord[-1]:
					redWord = paragraph.add_run('äu')
					redWord.font.color.rgb = RGBColor(255,0,0) #color red every noun word
				elif part == splitWord[0]:
					paragraph.add_run(part)

				else:
					paragraph.add_run(part)
		elif 'ee' in word:
			i = 0
			print(word.split('ee'))
			splitWord = word.split('ee')
			print(splitWord[-1])
			for part in splitWord:
				if part == splitWord[-1]:
					paragraph.add_run(part + ' ')

				elif part != splitWord[-1]:
					redWord = paragraph.add_run('ee')
					redWord.font.color.rgb = RGBColor(255,0,0) #color red every noun word
				elif part == splitWord[0]:
					paragraph.add_run(part)

				else:
					paragraph.add_run(part)
		elif 'eh' in word:
			i = 0
			print(word.split('eh'))
			splitWord = word.split('eh')
			print(splitWord[-1])
			for part in splitWord:
				if part == splitWord[-1]:
					paragraph.add_run(part + ' ')

				elif part != splitWord[-1]:
					redWord = paragraph.add_run('eh')
					redWord.font.color.rgb = RGBColor(255,0,0) #color red every noun word
				elif part == splitWord[0]:
					paragraph.add_run(part)
				else:
					paragraph.add_run(part)
		elif 'das' in word:
			redWord = paragraph.add_run('das ')
			redWord.font.color.rgb = RGBColor(255,0,0) #color red every noun word
		else:
			paragraph.add_run(word + ' ')	


	doc.add_page_break()	
	paragraph = doc.add_paragraph()
	run = paragraph.add_run('Aufgabe: Unterstreiche die richtigen Wortteile! (äu/eu), (eh/ee)), (dass/das)')	
	font = run.font
	font.size = Pt(18)
	font.bold = True			
	docnlp = nlp(' '.join(fullrawText))
	sents = docnlp.sents
	paragraph = doc.add_paragraph()
	for s in sents:
		sStr = str(s)					
		if 'eu' in sStr:					
			printS = sStr.replace('eu', '(äu/eu)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)
		elif 'äu' in sStr:					
			printS = sStr.replace('äu', '(äu/eu)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)

		elif 'ee' in sStr:
			printS = sStr.replace('ee', '(eh/ee)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)

		elif 'eh' in sStr:
			printS = sStr.replace('eh', '(eh/ee)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)

		elif 'das' in sStr:
			printS = sStr.replace('das', '(dass/das)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)

		elif 'dass' in sStr:
			printS = sStr.replace('dass', '(dass/das)')
			print(printS)
			run = paragraph.add_run(printS + '\n')
			font = run.font
			font.size = Pt(18)

	print('done')
	doc.save(os.path.join(os.path.expanduser('~'), 'microblog','pronounce-'+fileN))
	return 

