#! python3
# -*- coding: utf-8 -*-
import docx 	#python-docx module
from docx.shared import RGBColor
from docx.shared import Pt
import os

def docx_print(printText="",
                color=None,
                Paragraph="newParagraph",
                Bold=False,
                pageBreak=False,
                Doc=None,
                Table=None,
                NewTable=False,
                TableRow=0,
                TableCol=0,
                TableROWS=0,
                TableCOLS=0,
				save='' ):
  doc = Doc
  if save != '':
      save_path = os.path.join(os.path.expanduser('~'), save + 'fileTitle.docx')
      print("save")
      return save_path
  if pageBreak == True:
    doc.add_page_break()

  if Paragraph == "newParagraph":
    paragraph = doc.add_paragraph()
  elif Paragraph == "no":
    #print("no")
    pass
  else:
    paragraph = Paragraph

  if NewTable == True:

    table = doc.add_table(TableROWS ,TableCOLS)
    table.style = 'Table Grid'
  if Table != None:
    table1 = Table
    cell = table1.cell(TableRow,TableCol)
    cell.text = printText
    #table = doc.add_table(ROWS ,COLS)
    #table.style = 'Table Grid'




  if Paragraph != "no":
      run = paragraph.add_run(f'{printText}')
      font = run.font
      font.size = Pt(18)
      font.bold = Bold
      if color == "red":
        run.font.color.rgb = RGBColor(255,0,0) #color red every noun word
      return paragraph
  if NewTable == True:
      return table

def print_present_or_past(looked_up_verbs_list, found_verbs_list):
	import random
	import os

	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt

	#looked_up_verbs_list = []
	mix_list = []

	for elements in looked_up_verbs_list:
		for word in elements:
			if len(mix_list) == 26:
				break
			if word not in mix_list:
				mix_list.append(word)
	mix_list = list(set(mix_list))
	print('present or past mied list: ', mix_list )

	#print everythin to .docx file
	doc = docx.Document()

	save_path = docx_print(Doc= doc, save= 'presentorpast')

	paragraph = doc.add_paragraph()
	run = paragraph.add_run('Aufgabe: Hier siehst du Verben aus dem Text.\n \n  Ordne Sie richtig zu! \n Sind sie im Präsens oder im Präteritum?')
	font = run.font
	font.size = Pt(18)
	font.bold = True

	paragraph = doc.add_paragraph('')
	random.shuffle(mix_list)
	for word in mix_list:
		if word == mix_list[-1]:
			run = paragraph.add_run(word + ' ')
			font = run.font
			font.size = Pt(18)
			break
		run = paragraph.add_run(word + ', ')
		font = run.font
		font.size = Pt(18)


	table = doc.add_table(13 +1 ,2)
	table.style = 'Table Grid'

	cell = table.cell(0,0)
	cell.text = 'Präsens'
	paragraph = cell.add_paragraph()
	run = paragraph.add_run()
	font = run.font
	font.size = Pt(18)

	cell = table.cell(0,1)
	cell.text = 'Präteritum'
	#
	#save the result
	#doc.save(os.path.join(os.path.expanduser('~'), 'python-project', "kivy-test", "learny", "presentorpast" +"fileTitle.docx"))
	doc.save(save_path)

	return
