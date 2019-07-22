#! python3



def infinitive(inputtext, language):
	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt
	import os
	import spacy
	#from spacy.lang.de.examples import sentences
	from learny import docxprint, languageload
	#from flask import request
	#welcome message
	print('-------------------GENERATE infinit form of verbs docx has to be in a table-----------aim: ')

	# create document object
	doc = docx.Document()

	# variables and lists used
	i = 0
	verb_list = []
	some_verbs = []
	new_text = []
	infinit_verb_list = []
	textverb = []
	keptinfinitive = []
	Aufgabe = {
		"Kopfzeile": "Name: 				Klasse: 				Datum:  \n ",
		"Titel": "",
		"1. Aufgabe": "Aufgabe: Hier sind die Grundformen von Verben aus dem Text.\n Unterstreiche die gebeugte Form aus dem Text.\n",
		"Hinweise": "Hier sind die Wörter aus dem Text: \n Markiere die Wörter blau, die gleich geblieben sind. \n",
		"Rätselwörter": "Hier ein paar Rätselwörter aus dem Text: \n",
	}
	nlp = languageload.language_load(language)
	docnlp = nlp(inputtext) #load to spacy
	#nlp = spacy.load(request.form['language_options'])
	doc = docx.Document()# initializing python-docx
	save_path = docxprint.docx_print(Doc= doc, save= 'infinitive')

	#rawText = request.form['text']

	#docnlp = nlp(rawText)
	print(docnlp.text)
	paragraph = doc.add_paragraph()
	paragraph.text = ''
	#for token in docnlp:
	#	print(token.text, token.pos_, token.dep_)
	#print([t.lemma_ for t in doc])               # Lemmata der einzelnen Tokens
		#print(doc[1].lemma_)
		#print(t)
	for token in docnlp:
		if 'AUX' in token.pos_:
			print(token.lemma_)
			infinit_verb_list.append(token.lemma_)
			textverb.append(token.text)
			docxprint.docx_print(printText=str(token) + ' ', color="red", Paragraph=paragraph, Doc=doc)

		if 'VERB' in token.pos_:
			print(token.lemma_)
			infinit_verb_list.append(token.lemma_)
			textverb.append(token.text)

			docxprint.docx_print(printText=str(token) + ' ', color="red", Paragraph=paragraph, Doc=doc)
		elif 'AUX' in token.pos_:
			paragraph.add_run('')
		else:
			docxprint.docx_print(printText=str(token) + ' ', Paragraph=paragraph, Doc=doc)

	infinit_verb_list = list(set(infinit_verb_list))
	textverb = list(set(textverb))
	doc.add_page_break()
	#paragraph = doc.add_paragraph()
	docxprint.docx_print(printText=Aufgabe["1. Aufgabe"], Bold=True, Doc=doc)
	docxprint.docx_print(printText=Aufgabe["Hinweise"], Doc=doc)
	paragraph = doc.add_paragraph('')
	#random.shuffle(mix_list)
	for word in textverb:
		if word == textverb[-1]:
			docxprint.docx_print(printText=word + ' ', Paragraph=paragraph, Doc=doc)
			break
		docxprint.docx_print(printText=word + ', ', Paragraph=paragraph, Doc=doc)


	print("keptinfinitive:", keptinfinitive )
	#random.shuffle(verb_list)
	paragraph = doc.add_paragraph()
	for word in infinit_verb_list:
		docxprint.docx_print(printText=word + '_'*31+"\n", Paragraph=paragraph, Doc=doc)
	print('done')
	doc.save(save_path)
	return
