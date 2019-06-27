#! python3
# aim: help students learning to read by coloring red every second syllaba
# idea was seen in non free educational program milbenberger silbengenerator
# what the program does:
# reads an existing docx file with tables
# in words with more than one syllaba colors every second syllaba red
# creates a docx file to print the result
# TODO: create a document that keeps original formatting on run level i.e. with pictures and special formatting but with colored syllabas

def generateSylabas(inputtext):
	import subprocess
	import os
	#python-docx module code: pip install python-docx
	import docx
	from docx.shared import RGBColor

	#code: pip install PyHyphen
	from hyphen import Hyphenator
	de_DE = Hyphenator('de_DE')

	print('generate sylabas')
	print('input Text: ' + inputtext)
	# create document object
	doc = docx.Document()

	# variables and lists used
	i = 0
	word_print = []

	#get the input text from kivy
	text = inputtext
	#only work on text with '' specific words
	if text.find('') != -1:
		hyph_print = []
		word_print =[]
		hyph =[]

	#seperate text into words and put them in list hyph_input
		hyph_input = text.split()
		paragraph = doc.add_paragraph()
		paragraph.text = ''

	#work on separated words
		for w in hyph_input:
			hyph_print = []
			i = 1
			hyph = de_DE.syllables(w)
			if hyph == []:
				paragraph.add_run(w+' ')
			for syl in hyph:
				if i%2 == 0:
					redSyl = paragraph.add_run(syl)
					redSyl.font.color.rgb = RGBColor(255,0,0) #color red every second syllaba
				else :
					paragraph.add_run(syl)
				i +=1
			paragraph.add_run(' ')

	#save the result in absolute path
	doc.save(os.path.join(os.path.expanduser('~'), 'python-project','kivy-test', 'learny','syl'+ 'fileTitle.docx'))
	return
