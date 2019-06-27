#! python3
import docx 	#python-docx module
from docx.shared import RGBColor
from docx.shared import Pt

def docx_print(printText="",
                color=None,
                Paragraph="newParagraph",
                Bold=False,
                pageBreak=False,
                Doc=None,
                Table=None,
                NewTable=False,
                TableRow=0,
                TableCol=0,
                TableROWS=0,
                TableCOLS=0, ):
  doc = Doc
  if pageBreak == True:
    doc.add_page_break()

  if Paragraph == "newParagraph":
    paragraph = doc.add_paragraph()
  elif Paragraph == "no":
    #print("no")
    pass
  else:
    paragraph = Paragraph

  if NewTable == True:
    print("newTable")
    table = doc.add_table(TableROWS ,TableCOLS)
    table.style = 'Table Grid'
  if Table != None:
    table1 = Table
    cell = table1.cell(TableRow,TableCol)
    cell.text = printText
    #table = doc.add_table(ROWS ,COLS)
    #table.style = 'Table Grid'




  if Paragraph != "no":
      run = paragraph.add_run(f'{printText}')
      font = run.font
      font.size = Pt(18)
      font.bold = Bold
      if color == "red":
        run.font.color.rgb = RGBColor(255,0,0) #color red every noun word
      return paragraph
  if NewTable == True:
      return table
