#! python3
# -*- coding: utf-8 -*-
from __future__ import absolute_import, unicode_literals

def lookupDEMorphy(v):
	#look up 1person present tense and past tense of a word in the given text

	from demorphy import Analyzer
	import subprocess
	import os
	import shelve
	from hyphen import Hyphenator

	def search_known_verbs(iVerbs, Tense):

		if (iVerbs[1][0]['LEMMA'] == v.lemma_
			and iVerbs[1][0]['TENSE'] == Tense
			and iVerbs[1][0]['PERSON'] == '1per'
			and iVerbs[1][0]['NUMERUS'] == 'sing'
			and iVerbs[1][0]['MODE'] == 'ind'):
			return True
		else:
			return False

	def search_unknown_verbs(found, word, Tense, important_verb_list, printString):
		if (word[1].lemma == v.lemma_
			and word[1].tense == Tense
			and word[1].person == '1per'
			and word[1].numerus == 'sing'
			and word[1].mode == 'ind'):

			print('		Präsens: ich ' + word[0])
			#print(word)
			found += 1
			pres1 = 'ich ' + word[0]
			printString['pres1'] = pres1

			important_verb =[word[0], [{'LEMMA': word[1].lemma,
										'MODE': word[1].mode,
										'NUMERUS': word[1].numerus,
										'PERSON': word[1].person,
										'TENSE': word[1].tense}]]
			alreadyin = False
			for iVerbs in important_verb_list:
				#print(important_verb_list)
				if important_verb[0] == iVerbs[0] and word[1].mode == iVerbs[1][0]['MODE']:
					print('{Tense}1: already in important verb list: ' + iVerbs[0])
					alreadyin = True
				elif alreadyin == False and iVerbs[0] == important_verb_list[-1][0]:
					print('					new word saved: '+ str(important_verb))
					important_verb_list.append(important_verb)
					print('append')
					#print('important verb list: ' + str(important_verb_list[0][0]))
		return important_verb_list, printString, found


	"""shelfFile = shelve.open(os.path.join(
								os.path.expanduser('~'),
								'python-project' ,
								'kivy-test',
								'learny',
								'dictdatakivy'))"""
	#uncomment to create new shelf file database
	"""important_verb_list = [
								['aufbaue',
									[{'LEMMA': 'aufbauen',
									 'MODE': 'ind',
									 'NUMERUS': 'sing',
									 'PERSON': '1per',
									 'TENSE': 'pres'}]
								],
								['gebe',
									[{'LEMMA': 'geben',
									 'MODE': 'ind',
									 'NUMERUS': 'sing',
									 'PERSON': '1per',
									 'TENSE': 'pres'}]
								]
							]"""
	print('get important verb list')
	global important_verb_list
	important_verb_list = shelfFile['important_verb_list']
	important_verb = []
	printString = {}
	de_DE = Hyphenator('de_DE')
	firstHyph = ''
	letter_v = ''

	#create look up Terms
	#print(len(de_DE.syllables('bin')))
	if de_DE.syllables('bin') == []:
		print('Hyphenator: syllables returns empty')
	vStr = str(v).strip(' ')
	#first letter
	letter_v = vStr[0]
	#first syllable
	allHyph = de_DE.syllables(vStr)
	if allHyph == []:
		firstHyph = vStr
	else:
		firstHyph = allHyph[0]

	lookupProcess = [v.lemma_, firstHyph, 'ge', letter_v]
	#looking = iter(lookupProcess)
	found = 0
	for lookupTerm in lookupProcess:
		iVfound = 0

		if found == 3:
			continue


	# look up already known words
		print('looking up important_verb_list with:' + lookupTerm)
		for iVerbs in important_verb_list:

			# 1 Person Präsens singular
			if search_known_verbs(iVerbs, "pres"):
				pres1 = 'ich ' + iVerbs[0]
				print('	will be printed: Präsens: ich ' + iVerbs[0])
				printString['pres1'] = pres1
				iVfound += 1

			# 1 Person Präteritum singular
			if search_known_verbs(iVerbs, "past"):
				print('will be printed: Präteritum ich ' + iVerbs[0])
				past1 = 'ich ' + iVerbs[0]
				printString['past1'] = past1
				iVfound += 1


		#DEMorphy look up
		analyzer = Analyzer(char_subs_allowed=True)
		DEMorphy = analyzer.iter_lexicon_formatted(prefix=lookupTerm)
		#print(f"iVfound: {iVfound}")
		if iVfound <= 1:
			print('looking up DEMorphy with:' + lookupTerm + "\n"*5)

			for word in DEMorphy:
				if word == '':
					continue
				#Präsens sing

				important_verb_list, printString, found = search_unknown_verbs(found=found,
																		word=word,
																		Tense="pres",
																		printString=printString,
																		important_verb_list=important_verb_list)

				#Präteritum sing
				important_verb_list, printString, found = search_unknown_verbs(found=found,
																		word=word,
																		Tense="past",
																		printString=printString,
																		important_verb_list=important_verb_list)


	#print('one moment bevore putting new words into shelffILE')
	#shelfFile['important_verb_list'] = important_verb_list
	#shelfFile.close()
	#print('shelfFile closed')
	return printString

def flexTable(inputtext, langspinner):
	#from flask import request
	import subprocess
	import os
	#python-docx module
	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt
	import spacy
	from spacy.lang.de.examples import sentences
	import random
	from hyphen import Hyphenator
	import shelve

	print('generate flexion table'+ "\n"*5)


	# variables and lists used
	i = 0
	word_print = []
	verb_list = []
	some_verbs = []
	new_text = []
	infinit_verb_list = []
	new_list = []
	mix_list = []
	de_DE = Hyphenator('de_DE')
	global shelfFile
	shelfFile = shelve.open(os.path.join(
								os.path.expanduser('~'),
								'python-project' ,
								'kivy-test',
								'learny',
								'dictdatakivy'))
	# load language from kivy
	lang = ""
	if langspinner == "Sprache: Deutsch":
		lang = "de_core_news_sm"
	elif langspinner == "Sprache: Englisch":
		lang = "en_core_web_sm"
	else:
		print("didnt choose language"+ "\n"*5)
	nlp = spacy.load(lang)

	# load text from kivy to spacy
	docnlp = nlp(inputtext)
	print(docnlp.text)

	#creating list of verbs
	for token in docnlp:
		if 'AUX' in token.pos_:
			verb_list.append(token)

		if 'VERB' in token.pos_:
			verb_list.append(token)
	print(verb_list)

	#look up 1person present tense and past tense of a word in the given text
	for v in verb_list[0:-1]:
		print('beginning with look up of this word: ' + str(v)+ "\n"*5)
		vMorphy = lookupDEMorphy(v)
		new_list.append(list(vMorphy.values()))
		print('new_list result from look up' + str(new_list)+ "\n"*5)
		#print('ready with look up\n\n\')

	random.shuffle(new_list)
	for elements in new_list:
		for word in elements:
			mix_list.append(word)

	#print everythin to .docx file
	doc = docx.Document()
	paragraph = doc.add_paragraph()
	run = paragraph.add_run('Aufgabe: Hier siehst du Verben aus dem Text.\n \n  Ordne Sie richtig zu! \n Sind sie im Präsens oder im Präteritum?')
	font = run.font
	font.size = Pt(18)
	font.bold = True

	paragraph = doc.add_paragraph('')
	random.shuffle(mix_list)
	for word in mix_list:
		run = paragraph.add_run(word + ', ')
		font = run.font
		font.size = Pt(18)
		if word == mix_list[-1]:
			run = paragraph.add_run(word + ' ')

	table = doc.add_table(len(verb_list) +1 ,2)
	table.style = 'Table Grid'

	cell = table.cell(0,0)
	cell.text = 'Präsens'
	paragraph = cell.add_paragraph()
	run = paragraph.add_run()
	font = run.font
	font.size = Pt(18)

	cell = table.cell(0,1)
	cell.text = 'Präteritum'
	#
	#save the result
	doc.save(os.path.join(os.path.expanduser('~'), 'python-project', "kivy-test", "learny", 'flexTable'+"fileTitle.docx"))
	print('one moment bevore putting new words into shelffILE'+ "\n"*5)
	shelfFile['important_verb_list'] = important_verb_list
	shelfFile.close()
	print('shelfFile closed')

	return

text="""Firmen und Forschungseinrichtungen, die Algorithmen
	für Künstliche Intelligenz (KI) rennen, werden immer
	datenhungriger. Dies verdeutlichte Sophie Searcy, Data
	Scientist beim "KI-Bootcamp" Metis in New York, auf der
	am Samstag zu Ende gegangenen Konferenz "AI Traps" in
	Berlin. "Die Leute in Unternehmen reden nur noch darüber,
	 wie sie Daten bekommen." Mit immer umfangreicheren
	 Trainingsdaten-Sets wollten sie immer bessere
	 KI-Lösungsmodelle entwickeln."""
text2=""" Leistungsfähigere Algorithmen an sich seien nicht schlecht,
 	erläuterte Searcy. Sie hälfen dabei, dass Maschinen die "echte Welt" besser
	und schneller einschätzen könnten. Jedes KI-Modell sei letztlich eine kleine
	 Funktion, um "etwas Größeres einzufangen". Es gehe dabei um Lernprozesse in
	  der Form, dass Erfahrungen verarbeitet würden, um bestehende Modelle so zu
	   aktualisieren, dass sie auch für künftige Entwicklungen nützlich sind.
"""
flexTable(inputtext=text2, langspinner="Sprache: Deutsch")
