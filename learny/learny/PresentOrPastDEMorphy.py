#! python3
# -*- coding: utf-8 -*-
from __future__ import absolute_import, unicode_literals

def lookupDEMorphy(found_verb):
	#look up 1person present tense and past tense of a word in the given text

	from demorphy import Analyzer
	import subprocess
	import os
	import shelve
	from searchVerbs import search_verbs
	"""
	def search_known_verbs(knownVerbs, Tense):

		if (knownVerbs[1][0]['LEMMA'] == found_verb.lemma_
			and knownVerbs[1][0]['TENSE'] == Tense
			and knownVerbs[1][0]['PERSON'] == '1per'
			and knownVerbs[1][0]['NUMERUS'] == 'sing'
			and knownVerbs[1][0]['MODE'] == 'ind'):
			return True
		else:
			return False

	def search_unknown_verbs(foundInDEMorphy, word, Tense, known_verb_list, printString):
		if (word[1].lemma == found_verb.lemma_
			and word[1].tense == Tense
			and word[1].person == '1per'
			and word[1].numerus == 'sing'
			and word[1].mode == 'ind'):

			print(f'		{Tense}: ich ' + word[0])
			#print(word)
			#foundInDEMorphy += 1
			pres1 = 'ich ' + word[0]
			printString[f'{Tense}1'] = pres1

			newResult =[word[0], [{'LEMMA': word[1].lemma,
										'MODE': word[1].mode,
										'NUMERUS': word[1].numerus,
										'PERSON': word[1].person,
										'TENSE': word[1].tense}]]
			alreadyin = False
			for knownVerbs in known_verb_list:
				#print(known_verb_list)
				if newResult[0] == knownVerbs[0] and word[1].mode == knownVerbs[1][0]['MODE']:
					print(f'{Tense}1: already in known verb list: ' + knownVerbs[0])
					alreadyin = True
				elif alreadyin == False and knownVerbs[0] == known_verb_list[-1][0]:
					print('					new word saved: '+ str(newResult))
					known_verb_list.append(newResult)
					print('append')
					#print('important verb list: ' + str(known_verb_list[0][0]))
		return known_verb_list, printString, foundInDEMorphy"""



	#uncomment to create new shelf file database
	"""known_verb_list = [
								['aufbaue',
									[{'LEMMA': 'aufbauen',
									 'MODE': 'ind',
									 'NUMERUS': 'sing',
									 'PERSON': '1per',
									 'TENSE': 'pres'}]
								],
								['gebe',
									[{'LEMMA': 'geben',
									 'MODE': 'ind',
									 'NUMERUS': 'sing',
									 'PERSON': '1per',
									 'TENSE': 'pres'}]
								]
							]"""
	global known_verb_list
	newResult = []
	printString = {}
	firstHyph = ''
	firstLetter = ''

	#create look up Terms for lookupProcess
	preparedFoundVerb = str(found_verb).strip(' ')
	firstLetter = preparedFoundVerb[0]
	allHyph = de_DE.syllables(preparedFoundVerb)
	if allHyph == []:
		firstHyph = preparedFoundVerb
	else:
		firstHyph = allHyph[0]
	lookupProcess = [found_verb.lemma_, firstHyph, 'ge', firstLetter]

	foundInDEMorphy = 0
	knownVerbFound = ["",""]
	for lookupTerm in lookupProcess:
		print(f"foundInDEMorphy: {foundInDEMorphy}")
		if foundInDEMorphy == 3:
			print("will continue because foundIn DEMorphy is more than necessary")
			continue


	# look up already known words
		print('looking up known_verb_list with:\n' + "\t"*5 + lookupTerm + "\n")
		for knownVerbs in known_verb_list:

			# 1 Person Präsens singular
			if search_verbs.search_known_verbs(knownVerbs, "pres", found_verb):
				pres1 = 'ich ' + knownVerbs[0]
				print('\twill be printed: Präsens: \t ich ' + knownVerbs[0])
				printString['pres1'] = pres1
				knownVerbFound[0] = "pres found"


			# 1 Person Präteritum singular
			if search_verbs.search_known_verbs(knownVerbs, "past", found_verb):
				print('\twill be printed: Präteritum: \t ich ' + knownVerbs[0])
				past1 = 'ich ' + knownVerbs[0]
				printString['past1'] = past1
				knownVerbFound[1] = "past found"

		if knownVerbFound[0] == "pres found" and knownVerbFound[1] == "past found":
			print(f"already found everything. continue with next verb")
			break

		#DEMorphy look up
		analyzer = Analyzer(char_subs_allowed=True)
		DEMorphy = analyzer.iter_lexicon_formatted(prefix=lookupTerm)
		#print(f"knownVerbFound: {knownVerbFound}")
		if True:
			print('looking up \033[31mDEMorphy\033[0m with:' + lookupTerm + "\n"*5)
			for word in DEMorphy:
				if word == '':
					print("will continue word")
					continue
				#Präsens sing

				known_verb_list, printString, foundInDEMorphy = search_verbs.search_unknown_verbs(foundInDEMorphy=foundInDEMorphy,
																		word=word,
																		Tense="pres",
																		printString=printString,
																		known_verb_list=known_verb_list,
																		found_verb=found_verb)

				#Präteritum sing
				known_verb_list, printString, foundInDEMorphy = search_verbs.search_unknown_verbs(foundInDEMorphy=foundInDEMorphy,
																		word=word,
																		Tense="past",
																		printString=printString,
																		known_verb_list=known_verb_list,
																		found_verb=found_verb)


	#print('one moment bevore putting new words into shelffILE')
	#shelfFile['known_verb_list'] = known_verb_list
	#shelfFile.close()
	#print('shelfFile closed')
	return printString

def flexTable(inputtext, langspinner):
	#from flask import request
	import subprocess
	import os
	#python-docx module
	import docx
	from docx.shared import RGBColor
	from docx.shared import Pt
	import spacy
	from spacy.lang.de.examples import sentences
	import random
	from hyphen import Hyphenator
	import shelve

	print('generate flexion table'+ "\n"*5)


	# variables and lists used
	i = 0
	word_print = []
	found_verbs_list = []
	some_verbs = []
	new_text = []
	infinit_verb_list = []
	new_list = []
	mix_list = []
	#knownVerbFound = 0
	global de_DE
	de_DE = Hyphenator('de_DE')
	global shelfFile
	shelfFile = shelve.open(os.path.join(
								os.path.expanduser('~'),
								'python-project' ,
								'kivy-test',
								'learny',
								'dictdatakivy'))
	print('get known verb list')
	global known_verb_list
	known_verb_list = shelfFile['important_verb_list']
	#load language from kivy
	lang = ""
	if langspinner == "Sprache: Deutsch":
		lang = "de_core_news_sm"
	elif langspinner == "Sprache: Englisch":
		lang = "en_core_web_sm"
	else:
		print("didnt choose language"+ "\n"*5)
	nlp = spacy.load(lang)

	# load text from kivy to spacy
	docnlp = nlp(inputtext)
	print(docnlp.text)

	#creating list of verbs
	for token in docnlp:
		if 'AUX' in token.pos_:
			found_verbs_list.append(token)

		if 'VERB' in token.pos_:
			found_verbs_list.append(token)
	print(found_verbs_list)

	#look up 1person present tense and past tense of a word in the given text
	for found_verb in found_verbs_list[0:-1]:
		print("\n"*5 + '\033[31mbeginning\033[0m with look up of this word: ' + str(found_verb)+ "\n"*2)
		vMorphy = lookupDEMorphy(found_verb)
		new_list.append(list(vMorphy.values()))
		#print("\n"*5 + 'new_list result from look up\n' + str(new_list)+ "\n"*5)
		print("\n"*2 + 'new_list result from look up: ')
		for result in new_list:
			print(result)
		print("\n"*5)
		#print('ready with look up\n\n\')

	random.shuffle(new_list)
	for elements in new_list:
		for word in elements:
			mix_list.append(word)

	#print everythin to .docx file
	doc = docx.Document()
	paragraph = doc.add_paragraph()
	run = paragraph.add_run('Aufgabe: Hier siehst du Verben aus dem Text.\n \n  Ordne Sie richtig zu! \n Sind sie im Präsens oder im Präteritum?')
	font = run.font
	font.size = Pt(18)
	font.bold = True

	paragraph = doc.add_paragraph('')
	random.shuffle(mix_list)
	for word in mix_list:
		run = paragraph.add_run(word + ', ')
		font = run.font
		font.size = Pt(18)
		if word == mix_list[-1]:
			run = paragraph.add_run(word + ' ')

	table = doc.add_table(len(found_verbs_list) +1 ,2)
	table.style = 'Table Grid'

	cell = table.cell(0,0)
	cell.text = 'Präsens'
	paragraph = cell.add_paragraph()
	run = paragraph.add_run()
	font = run.font
	font.size = Pt(18)

	cell = table.cell(0,1)
	cell.text = 'Präteritum'
	#
	#save the result
	doc.save(os.path.join(os.path.expanduser('~'), 'python-project', "kivy-test", "learny", 'flexTable'+"fileTitle.docx"))
	print('one moment bevore putting new words into shelffILE')
	#comment because of tests
	#shelfFile['important_verb_list'] = known_verb_list
	shelfFile.close()
	print('shelfFile closed')

	return

text="""Firmen und Forschungseinrichtungen, die Algorithmen
	für Künstliche Intelligenz (KI) rennen, werden immer
	datenhungriger. Dies verdeutlichte Sophie Searcy, Data
	Scientist beim "KI-Bootcamp" Metis in New York, auf der
	am Samstag zu Ende gegangenen Konferenz "AI Traps" in
	Berlin. "Die Leute in Unternehmen reden nur noch darüber,
	 wie sie Daten bekommen." Mit immer umfangreicheren
	 Trainingsdaten-Sets wollten sie immer bessere
	 KI-Lösungsmodelle entwickeln."""
text2=""" Leistungsfähigere Algorithmen an sich kenne nicht schlecht,
 	rennen Searcy. Sie hälfen dabei, dass Maschinen die "echte Welt" besser
	und schneller einschätzen könnten. Jedes KI-Modell gehe letztlich eine kleine
	 Funktion, um "etwas Größeres einzufangen". Es gehe dabei um Lernprozesse in
	  der Form, dass Erfahrungen verarbeitet würden, um bestehende Modelle so zu
	   aktualisieren, dass sie auch für künftige Entwicklungen nützlich fallen.
"""
flexTable(inputtext=text2, langspinner="Sprache: Deutsch")
